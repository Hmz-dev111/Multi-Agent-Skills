# skills/knowledge_report/knowledge_report.py
"""
knowledge_report 技能的代码实现
包含：
  - ChromaDB 向量检索（供 KnowledgeAgent retrieve 阶段使用）
  - DOCX / JSON 报告生成（供 KnowledgeAgent report 阶段使用）
"""
import os
import json
from datetime import datetime
from typing import List, Dict, Any

# ── 向量检索 ──────────────────────────────────────────────────────────────────

CHROMA_PATH = os.path.abspath(
    os.path.join(os.path.dirname(__file__), '..', '..', 'chroma_db')
)

_chroma_client     = None
_chroma_collection = None
_embedder          = None


def _get_retrieval_components():
    global _chroma_client, _chroma_collection, _embedder
    if _chroma_client is None:
        import chromadb
        from sentence_transformers import SentenceTransformer
        _chroma_client     = chromadb.PersistentClient(path=CHROMA_PATH)
        _chroma_collection = _chroma_client.get_or_create_collection(
            name="cvrptw_docs",
            metadata={"hnsw:space": "cosine"}
        )
        _embedder = SentenceTransformer("/root/.cache/huggingface/hub/models--BAAI--bge-small-zh-v1.5/snapshots/7999e1d3359715c523056ef9478215996d62a620")
    return _chroma_collection, _embedder


def retrieve(query: str, top_k: int = 5) -> List[Dict[str, Any]]:
    """检索本地知识库，返回相关文档列表"""
    collection, embedder = _get_retrieval_components()
    if collection.count() == 0:
        return []

    embedding = embedder.encode([query]).tolist()
    results   = collection.query(
        query_embeddings=embedding,
        n_results=min(top_k, collection.count())
    )

    docs = []
    for doc, meta, dist in zip(
        results['documents'][0],
        results['metadatas'][0],
        results['distances'][0]
    ):
        docs.append({
            "content":         doc,
            "source":          meta.get("source", "unknown"),
            "relevance_score": round(1 - dist, 3)
        })
    return docs


def format_for_prompt(docs: List[Dict[str, Any]]) -> str:
    """将检索结果格式化为适合插入 prompt 的文本"""
    if not docs:
        return "（本地知识库暂无相关内容）"
    parts = []
    for i, d in enumerate(docs):
        parts.append(f"[{i+1}] 来源: {d['source']} (相关度: {d['relevance_score']})\n{d['content']}")
    return "\n\n---\n\n".join(parts)


# ── 报告生成 ──────────────────────────────────────────────────────────────────

def save_json(report_data: Dict[str, Any], output_dir: str = "output") -> str:
    """保存 JSON 格式报告"""
    os.makedirs(output_dir, exist_ok=True)
    ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = os.path.join(output_dir, f"cvrptw_report_{ts}.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(report_data, f, ensure_ascii=False, indent=2)
    return path


def save_docx(report_data: Dict[str, Any], output_dir: str = "output") -> str:
    """保存 DOCX 格式报告"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("请先安装 python-docx：pip install python-docx")

    os.makedirs(output_dir, exist_ok=True)
    doc = Document()

    # 标题
    title_para = doc.add_heading(report_data.get("title", "CVRPTW 求解分析报告"), level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"生成时间：{report_data.get('generated_at', '')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # 1. 问题概述
    doc.add_heading("1. 问题概述", level=1)
    ps = report_data.get("problem_summary", {})
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text, hdr[1].text = "属性", "值"
    for k, v in ps.items():
        row = tbl.add_row().cells
        row[0].text, row[1].text = str(k), str(v)
    doc.add_paragraph("")

    # 2. 求解过程
    doc.add_heading("2. 初始解生成", level=1)
    sp = report_data.get("solving_process", {})
    doc.add_paragraph(f"求解说明：{sp.get('reasoning', 'N/A')}")
    doc.add_paragraph(f"初始方案：{sp.get('initial_solution', [])}")
    doc.add_paragraph(f"初始成本：{sp.get('initial_cost', 'N/A')}")

    # 3. 验证优化
    doc.add_heading("3. 验证与优化结果", level=1)
    vr = report_data.get("validation_result", {})
    status_map = {"feasible": "✅ 可行", "infeasible": "❌ 不可行", "improved": "✅ 已优化"}
    doc.add_paragraph(f"状态：{status_map.get(vr.get('status',''), vr.get('status',''))}")
    doc.add_paragraph(f"优化后成本：{vr.get('final_cost', 'N/A')}")
    doc.add_paragraph(f"成本改进率：{vr.get('improvement_rate', 0)}%")
    for v in vr.get("violations", []):
        doc.add_paragraph(f"  • {v}", style="List Bullet")

    # 4. 路线详情
    route_details = report_data.get("route_details", [])
    if route_details:
        doc.add_heading("4. 路线详情", level=1)
        rt = doc.add_table(rows=1, cols=5)
        rt.style = "Table Grid"
        for i, h in enumerate(["路线", "客户序列", "站点数", "总距离", "载重利用率"]):
            rt.rows[0].cells[i].text = h
        for rd in route_details:
            row = rt.add_row().cells
            row[0].text = f"路线 {rd.get('route_id', '')}"
            row[1].text = str(rd.get("stops", []))
            row[2].text = str(rd.get("num_stops", ""))
            row[3].text = str(rd.get("total_distance", ""))
            row[4].text = f"{rd.get('load_utilization', '')}%"
        doc.add_paragraph("")

    # 5. 领域知识
    doc.add_heading("5. 领域知识对照", level=1)
    doc.add_paragraph(report_data.get("knowledge_context", "暂无"))
    doc.add_paragraph(f"算法建议：{report_data.get('algorithm_suggestion', '暂无')}")

    # 6. 结论
    doc.add_heading("6. 结论与建议", level=1)
    doc.add_paragraph(report_data.get("conclusion", "暂无"))

    ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = os.path.join(output_dir, f"cvrptw_report_{ts}.docx")
    doc.save(path)
    return path